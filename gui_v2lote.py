import re
import os
import glob
import pandas as pd
from docx import Document
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

# ==========================
# Utilidades
# ==========================

def extrair_dados_texto(texto: str) -> dict:
    """Extrai Nome, CPF e Data a partir de um texto plano."""
    nome = re.search(r"(?:Nome|Nome completo)[:\-–]?\s*(.+)", texto, re.IGNORECASE)
    cpf = re.search(r"\b\d{3}\.\d{3}\.\d{3}-\d{2}\b", texto)
    data = re.search(r"\b\d{2}[/-]\d{2}[/-]\d{4}\b", texto)

    return {
        "Nome": nome.group(1).strip() if nome else "",
        "CPF": cpf.group(0) if cpf else "",
        "Data": data.group(0) if data else "",
    }


def extrair_dados_docx(file_path: str) -> dict:
    """Lê um arquivo DOCX e tenta extrair dados dos parágrafos e tabelas."""
    doc = Document(file_path)

    # 1) Tenta nos parágrafos
    texto_paragrafos = "\n".join(p.text for p in doc.paragraphs)
    dados = extrair_dados_texto(texto_paragrafos)

    # 2) Se vazio, tenta nas tabelas
    if not any(dados.values()):
        for tabela in doc.tables:
            for linha in tabela.rows:
                linha_texto = " ".join(cell.text for cell in linha.cells)
                dados = extrair_dados_texto(linha_texto)
                if any(dados.values()):
                    break
            if any(dados.values()):
                break

    return dados


def extrair_em_lote(file_paths: list[str]) -> list[dict]:
    """Processa vários DOCX e retorna lista de dicionários com os campos + nome do arquivo."""
    resultados = []
    for path in file_paths:
        try:
            dados = extrair_dados_docx(path)
            base = os.path.basename(path)
            resultados.append({
                "Arquivo": base,
                "Nome": dados.get("Nome", ""),
                "CPF": dados.get("CPF", ""),
                "Data": dados.get("Data", ""),
            })
        except Exception as e:
            # Registra a falha como linha também (útil para auditoria)
            resultados.append({
                "Arquivo": os.path.basename(path),
                "Nome": "",
                "CPF": "",
                "Data": "",
                "Erro": str(e),
            })
    return resultados


def gerar_excel_em_arquivo_lote(linhas: list[dict], save_path: str) -> None:
    df = pd.DataFrame(linhas)
    # Ordena colunas de forma amigável, se existirem
    cols = [c for c in ["Arquivo", "Nome", "CPF", "Data", "Erro"] if c in df.columns]
    df = df[cols + [c for c in df.columns if c not in cols]]
    df.to_excel(save_path, index=False)


# ==========================
# Interface Tkinter
# ==========================

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Extrair DOCX para EXCEL (Lote)")
        self.geometry("780x460")
        self.minsize(720, 420)

        self.current_files = []    # lista de caminhos selecionados
        self.result_rows = []      # linhas extraídas

        self._build_ui()

    def _build_ui(self):
        # Top: seleção de arquivo/pasta
        top = ttk.Frame(self, padding=12)
        top.pack(fill=tk.X)

        ttk.Label(top, text="Arquivos DOCX:").pack(side=tk.LEFT)
        ttk.Button(top, text="Abrir… (múltiplos)", command=self.on_open_many)\
            .pack(side=tk.LEFT, padx=(8, 0))
        ttk.Button(top, text="Abrir pasta…", command=self.on_open_folder)\
            .pack(side=tk.LEFT, padx=(8, 0))

        # Opção de incluir subpastas
        self.recursive_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(top, text="Incluir subpastas", variable=self.recursive_var)\
            .pack(side=tk.LEFT, padx=12)

        # Centro: tabela de resultados
        mid = ttk.Frame(self, padding=(12, 6))
        mid.pack(fill=tk.BOTH, expand=True)

        cols = ("Arquivo", "Nome", "CPF", "Data")
        self.tree = ttk.Treeview(mid, columns=cols, show="headings", height=12)
        for c in cols:
            self.tree.heading(c, text=c)
            self.tree.column(c, width=160, anchor=tk.W)
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        vsb = ttk.Scrollbar(mid, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=vsb.set)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)

        # Rodapé: botões de ação
        actions = ttk.Frame(self, padding=12)
        actions.pack(fill=tk.X)
        ttk.Button(actions, text="Processar", command=self.on_process).pack(side=tk.LEFT)
        ttk.Button(actions, text="Salvar Excel", command=self.on_save_excel).pack(side=tk.LEFT, padx=8)
        ttk.Button(actions, text="Limpar", command=self.on_clear).pack(side=tk.LEFT)
        ttk.Button(actions, text="Remover seleção", command=self.on_remove_selected).pack(side=tk.LEFT, padx=8)

        self.status_var = tk.StringVar(value="Pronto.")
        ttk.Label(self, textvariable=self.status_var, anchor=tk.W).pack(fill=tk.X, padx=12, pady=(0,12))

    # ===== Ações =====
    def on_open_many(self):
        paths = filedialog.askopenfilenames(
            title="Selecionar arquivos DOCX",
            filetypes=[("Documentos do Word", "*.docx")],
        )
        if paths:
            adicionados = 0
            for p in paths:
                if p not in self.current_files:
                    self.current_files.append(p)
                    adicionados += 1
            self.status_var.set(f"{adicionados} arquivo(s) adicionado(s). Total: {len(self.current_files)}.")

    def on_open_folder(self):
        folder = filedialog.askdirectory(title="Selecionar pasta")
        if not folder:
            return
        pattern = "**/*.docx" if self.recursive_var.get() else "*.docx"
        encontrados = glob.glob(os.path.join(folder, pattern), recursive=self.recursive_var.get())
        if not encontrados:
            messagebox.showinfo("Info", "Nenhum arquivo .docx encontrado.")
            return
        adicionados = 0
        for p in encontrados:
            if p not in self.current_files:
                self.current_files.append(p)
                adicionados += 1
        self.status_var.set(f"{adicionados} arquivo(s) da pasta adicionados. Total: {len(self.current_files)}.")

    def on_process(self):
        if not self.current_files:
            messagebox.showwarning("Aviso", "Selecione arquivos ou uma pasta primeiro.")
            return
        self.status_var.set("Processando… isso pode levar alguns segundos em lotes maiores.")
        self.update_idletasks()

        self.result_rows = extrair_em_lote(self.current_files)
        # Popular a tabela
        for i in self.tree.get_children():
            self.tree.delete(i)
        for row in self.result_rows:
            self.tree.insert("", tk.END, values=(row.get("Arquivo",""), row.get("Nome",""), row.get("CPF",""), row.get("Data","")))

        ok = sum(1 for r in self.result_rows if any([r.get("Nome"), r.get("CPF"), r.get("Data")]))
        falhas = sum(1 for r in self.result_rows if r.get("Erro"))
        self.status_var.set(f"Processo concluído. {ok} com dados extraídos, {falhas} com erro, total {len(self.result_rows)}.")

    def on_save_excel(self):
        if not self.result_rows:
            messagebox.showwarning("Aviso", "Nada para salvar: processe arquivos antes.")
            return
        path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Planilha do Excel", "*.xlsx")],
            title="Salvar como",
            initialfile="dados_lote.xlsx",
        )
        if path:
            try:
                gerar_excel_em_arquivo_lote(self.result_rows, path)
                self.status_var.set(f"Excel salvo em: {path}")
            except Exception as e:
                messagebox.showerror("Erro", f"Não foi possível salvar o Excel.\n\n{e}")

    def on_remove_selected(self):
        """Remove da lista de processamento os arquivos selecionados na tabela (pelo nome)."""
        selecionados = self.tree.selection()
        if not selecionados:
            messagebox.showinfo("Info", "Selecione linha(s) na tabela para remover.")
            return
        nomes = {self.tree.item(i, "values")[0] for i in selecionados}  # Arquivo
        # Remove dos paths com base no nome do arquivo
        self.current_files = [p for p in self.current_files if os.path.basename(p) not in nomes]
        # Remove da visualização (apenas tabela de resultado)
        for i in selecionados:
            self.tree.delete(i)
        # E também das linhas processadas (se já houver)
        if self.result_rows:
            self.result_rows = [r for r in self.result_rows if r.get("Arquivo") not in nomes]
        self.status_var.set(f"Removidos {len(nomes)} arquivo(s). Restam {len(self.current_files)} para processar.")

    def on_clear(self):
        self.current_files = []
        self.result_rows = []
        for i in self.tree.get_children():
            self.tree.delete(i)
        self.status_var.set("Pronto.")


if __name__ == "__main__":
    # Estilo básico (Windows): melhora nitidez em telas de alta densidade
    try:
        from ctypes import windll  # type: ignore
        windll.shcore.SetProcessDpiAwareness(1)
    except Exception:
        pass

    app = App()
    app.mainloop()
