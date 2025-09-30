import re
import pandas as pd
from docx import Document
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

# ==========================
# Utilidades
# ==========================

def extrair_dados_texto(texto: str) -> dict:
    """Extrai Nome, CPF e Data a partir de um texto plano."""
    nome = re.search(r"(?:Nome|Nome completo)[:\-–]?\s*(.+)", texto, re.IGNORECASE)
    cpf = re.search(r"\b\d{3}\.\d{3}\.\d{3}-\d{2}\b", texto)
    data = re.search(r"\b\d{2}[/-]\d{2}[/-]\d{4}\b", texto)

    return {
        "Nome": nome.group(1).strip() if nome else "",
        "CPF": cpf.group(0) if cpf else "",
        "Data": data.group(0) if data else "",
    }


def extrair_dados_docx(file_path: str) -> dict:
    """Lê um arquivo DOCX e tenta extrair dados dos parágrafos e tabelas."""
    doc = Document(file_path)

    # 1) Tenta nos parágrafos
    texto_paragrafos = "\n".join(p.text for p in doc.paragraphs)
    dados = extrair_dados_texto(texto_paragrafos)

    # 2) Se vazio, tenta nas tabelas
    if not any(dados.values()):
        for tabela in doc.tables:
            for linha in tabela.rows:
                linha_texto = " ".join(cell.text for cell in linha.cells)
                dados = extrair_dados_texto(linha_texto)
                if any(dados.values()):
                    break
            if any(dados.values()):
                break

    return dados


def gerar_excel_em_arquivo(dados: dict, save_path: str) -> None:
    df = pd.DataFrame([dados])
    df.to_excel(save_path, index=False)


# ==========================
# Interface Tkinter
# ==========================

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Extrair DOCX para EXCEL")
        self.geometry("360x100")
        self.minsize(560, 150)

        self.current_file_path = None
        self.dados = {"Nome": "", "CPF": "", "Data": ""}

        self._build_ui()

    def _build_ui(self):
        # Top: seleção de arquivo
        top = ttk.Frame(self, padding=12)
        top.pack(fill=tk.X)

        self.path_var = tk.StringVar()
        ttk.Label(top, text="Arquivo DOCX:").pack(side=tk.LEFT)
        self.path_entry = ttk.Entry(top, textvariable=self.path_var)
        self.path_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=8)
        ttk.Button(top, text="Abrir...", command=self.on_open).pack(side=tk.LEFT)

     
        self.nome_var = tk.StringVar()
        self.cpf_var = tk.StringVar()
        self.data_var = tk.StringVar()

        

        # Rodapé: botões de ação
        actions = ttk.Frame(self, padding=12)
        actions.pack(fill=tk.X)
        center = ttk.Frame(actions)
        center.pack(expand=True)

        ttk.Button(center, text="Processar", command=self.on_process).pack(side=tk.LEFT, padx=8)
        ttk.Button(center, text="Salvar Excel", command=self.on_save_excel).pack(side=tk.LEFT, padx=8)
        ttk.Button(center, text="Limpar", command=self.on_clear).pack(side=tk.LEFT, padx=8)

        self.status_var = tk.StringVar(value="Pronto.")
        ttk.Label(self, textvariable=self.status_var, anchor="center").pack(fill=tk.X, padx=12, pady=(0,12))

    def _row(self, parent, r, label, var):
        ttk.Label(parent, text=f"{label}:").grid(row=r, column=0, sticky=tk.W, padx=(0,8), pady=6)
        e = ttk.Entry(parent, textvariable=var)
        e.grid(row=r, column=1, sticky=tk.EW, pady=6)
        parent.grid_columnconfigure(1, weight=1)

    # ===== Ações =====
    def on_open(self):
        path = filedialog.askopenfilename(
            title="Selecionar arquivo DOCX",
            filetypes=[("Documentos do Word", "*.docx")],
        )
        if path:
            self.current_file_path = path
            self.path_var.set(path)
            self.status_var.set("Arquivo selecionado. Clique em Processar.")

    def on_process(self):
        if not self.current_file_path:
            messagebox.showwarning("Aviso", "Selecione um arquivo DOCX primeiro.")
            return
        try:
            self.dados = extrair_dados_docx(self.current_file_path)
            self.nome_var.set(self.dados.get("Nome", ""))
            self.cpf_var.set(self.dados.get("CPF", ""))
            self.data_var.set(self.dados.get("Data", ""))
            if any(self.dados.values()):
                self.status_var.set("Dados extraídos com sucesso.")
            else:
                self.status_var.set("Nenhum dado correspondente encontrado com as regex atuais.")
        except Exception as e:
            messagebox.showerror("Erro", f"Falha ao processar o arquivo.\n\n{e}")

    def on_save_excel(self):
        if not any([self.nome_var.get(), self.cpf_var.get(), self.data_var.get()]):
            messagebox.showwarning("Aviso", "Nada para salvar: processe um arquivo antes.")
            return
        path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Planilha do Excel", "*.xlsx")],
            title="Salvar como",
            initialfile="dados.xlsx",
        )
        if path:
            try:
                gerar_excel_em_arquivo(
                    {"Nome": self.nome_var.get(), "CPF": self.cpf_var.get(), "Data": self.data_var.get()},
                    path,
                )
                self.status_var.set(f"Excel salvo em: {path}")
            except Exception as e:
                messagebox.showerror("Erro", f"Não foi possível salvar o Excel.\n\n{e}")

    def on_clear(self):
        self.current_file_path = None
        self.path_var.set("")
        self.nome_var.set("")
        self.cpf_var.set("")
        self.data_var.set("")
        self.status_var.set("Pronto.")


if __name__ == "__main__":
    # Estilo básico (Windows): melhora nitidez em telas de alta densidade
    try:
        from ctypes import windll  # type: ignore
        windll.shcore.SetProcessDpiAwareness(1)
    except Exception:
        pass

    app = App()
    app.mainloop()
