from docx import Document
import pandas as pd
import re

# === Funções auxiliares ===
def extrair_dados_texto(texto):
    # Regex para nome (após "Nome:" ou "nome completo:")
    nome = re.search(r"(?:Nome|Nome completo)[:\-–]?\s*(.+)", texto, re.IGNORECASE)
    
    # Regex para CPF
    cpf = re.search(r"\d{3}\.\d{3}\.\d{3}-\d{2}", texto)

    # Regex para data (ex: 01/01/2025 ou 01-01-2025)
    data = re.search(r"\d{2}[/-]\d{2}[/-]\d{4}", texto)

    return {
        "Nome": nome.group(1).strip() if nome else "",
        "CPF": cpf.group(0) if cpf else "",
        "Data": data.group(0) if data else ""
    }

# === Leitura do Word ===
doc = Document("word_teste_dados.docx")

# Junta o texto dos parágrafos
texto_paragrafos = "\n".join([p.text for p in doc.paragraphs])

# Tenta extrair dos parágrafos
dados = extrair_dados_texto(texto_paragrafos)

# Se estiver vazio, tenta nas tabelas
if not any(dados.values()):
    for tabela in doc.tables:
        for linha in tabela.rows:
            linha_texto = " ".join([cell.text for cell in linha.cells])
            dados = extrair_dados_texto(linha_texto)
            if any(dados.values()):
                break
        if any(dados.values()):
            break

# === Exporta para Excel ===
df = pd.DataFrame([dados])
df.to_excel("dados.xlsx", index=False)

print("✅ Dados extraídos e salvos em 'dados.xlsx':")
print(df)
